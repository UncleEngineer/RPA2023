import pytesseract
from docx import Document

def scanEng(imageFile):
    text = pytesseract.image_to_string(imageFile)
    return text

def scanThai(imageFile):
    text = pytesseract.image_to_string(imageFile, lang='tha')
    return text

textEng = scanEng('bookeng.jpg')
print(textEng)
print('===========================================')
textThai = scanThai('bookthai.jpg')
print(textThai)

def writeWordFile():
    document = Document()

    document.add_heading('Scan English text from OCR')
    p = document.add_paragraph(textEng)
    p.add_run('\n\nby Uncle Engineer')

    document.add_heading('การอ่านข้อความภาษาไทยจาก OCR')
    p = document.add_paragraph(textThai)
    p.add_run('\n\nโดย ลุงวิศวกร สอนคำนวณ')

    document.add_page_break()

    document.save('ocr_text.docx')

writeWordFile()

textform = pytesseract.image_to_string('formthai.png', lang='tha')
# print(textform)

raw_text = textform.split(' ')
#print(raw_text)

cleaned_text = []

for t in raw_text:
    if t != '':
        cleaned_text.append(t.strip().replace('\n',''))

# print(cleaned_text)

checktext = 'start'
company_name = ''

for s in cleaned_text:
    if 'ชื่อสถานประกอบการ' in checktext:
        company_name = s
        print(company_name)
    checktext = s